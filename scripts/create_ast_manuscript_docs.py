from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9D9D9", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)


def add_paragraph(doc, text="", style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_width(hdr[i], widths[i])
        set_cell_shading(hdr[i], GRAY_FILL)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def build_main_manuscript():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Correlation-informed multi-fidelity Bayesian optimization for film-cooling hole layout on a C3X turbine vane")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Anonymized manuscript skeleton for Aerospace Science and Technology")
    r.italic = True

    add_paragraph(doc, "Journal compliance note", style="Heading 1")
    add_paragraph(
        doc,
        "This file is the anonymized main manuscript skeleton. Do not add author names, affiliations, acknowledgements, grant identifiers that reveal identity, or self-identifying file names here. AST requires the title page with author details to be submitted separately."
    )
    add_paragraph(
        doc,
        "Source basis: AST Guide for Authors, ScienceDirect, accessed during drafting on 2026-06-24."
    )

    add_paragraph(doc, "Highlights", style="Heading 1")
    add_paragraph(doc, "Submit the final highlights as a separate editable file. Each bullet must be no more than 85 characters including spaces.")
    add_bullets(
        doc,
        [
            "Film-cooling correlations define a Gaussian-process prior mean.",
            "Multi-fidelity BO jointly selects C3X designs and CFD fidelity.",
            "Learnable prior coefficients improve robustness to correlation bias.",
            "Cost-aware acquisition reduces equivalent high-fidelity CFD calls.",
            "C3X validation links optimization gains to turbine-vane cooling.",
        ],
    )

    add_paragraph(doc, "Abstract", style="Heading 1")
    add_paragraph(
        doc,
        "[Draft after results are available; maximum 250 words.] Film-cooling layout optimization for turbine vanes requires repeated CFD evaluations, while conventional black-box Bayesian optimization does not exploit established film-cooling correlations. This work proposes a correlation-informed multi-fidelity Bayesian optimization framework in which classical row effectiveness correlations and Sellers superposition define the prior mean of a Gaussian-process surrogate. The surrogate learns the residual between CFD and the physics prior, with prior coefficients updated from data. The framework combines a zero-cost correlation model, coarse RANS, and fine RANS through a cost-aware acquisition function. The method is evaluated on a C3X turbine vane film-cooling case with automated CAD, meshing, solution, and post-processing. [RESULT: validation accuracy.] [RESULT: L1/L2 correlation.] [RESULT: equivalent L2 reduction.] [RESULT: final eta_bar improvement.] The results demonstrate [SUPPORTED CLAIM ONLY AFTER DATA]."
    )

    add_paragraph(doc, "Keywords", style="Heading 1")
    add_paragraph(doc, "Film cooling; Bayesian optimization; Multi-fidelity CFD; Gaussian process; Turbine vane; Physics-informed surrogate; C3X")

    add_paragraph(doc, "Nomenclature", style="Heading 1")
    add_table(
        doc,
        ["Symbol", "Definition", "Unit"],
        [
            ("eta_aw", "Adiabatic film-cooling effectiveness", "-"),
            ("eta_bar", "Area-averaged effectiveness over the protected surface", "-"),
            ("m_phys", "Physics-informed GP prior mean", "-"),
            ("M", "Blowing ratio", "-"),
            ("zeta", "Total pressure loss coefficient", "-"),
            ("lambda_l", "Measured cost of fidelity level l", "core-hour or wall-time"),
        ],
        [1600, 6100, 1660],
    )

    sections = [
        ("1. Introduction", [
            "Motivation: turbine-vane film cooling is an aerospace propulsion design problem with expensive CFD evaluations.",
            "Gap: black-box BO ignores film-cooling correlations and superposition models that already encode engineering knowledge.",
            "Contribution: correlation-informed GP prior mean, learnable prior coefficients, cost-aware multi-fidelity BO, C3X validation, and sample-efficiency evidence.",
        ]),
        ("2. Related work", [
            "2.1 Film-cooling design and optimization.",
            "2.2 Bayesian optimization for expensive CFD and aerospace design.",
            "2.3 Multi-fidelity surrogate modeling and acquisition functions.",
            "2.4 Prior-guided and physics-informed Bayesian optimization.",
        ]),
        ("3. Problem formulation", [
            "Define design vector x using nondimensional variables such as s/C, p/D, alpha, and optional beta.",
            "Define objective eta_bar and constraints on coolant mass flow and total pressure loss.",
            "Define paper-grid high-fidelity evaluation count and equivalent cost for fair comparison across L2 and L3 CFD evaluations.",
        ]),
        ("4. Correlation-informed Bayesian surrogate", [
            "Use row-level film-cooling effectiveness correlations and Sellers superposition to construct m_phys(x; theta).",
            "Model f(x) = m_phys(x; theta) + r(x), where r is a zero-mean GP residual.",
            "Estimate theta and kernel hyperparameters by MAP; compare fixed and learnable prior variants.",
        ]),
        ("5. Multi-fidelity optimization framework", [
            "Use L1 correlation knowledge, L2 coarse RANS, and L3 paper-quality RANS; reserve the fine mesh for external audits.",
            "Run paired L2/L3 CFD diagnostics before committing to co-Kriging or NARGP.",
            "Use cost-aware acquisition alpha(x,l)/lambda_l and probability-of-feasibility factors for constraints.",
        ]),
        ("6. C3X CFD case setup and validation", [
            "Describe C3X geometry, film-row layout, computational domain, mesh tiers, boundary conditions, and solver settings.",
            "Validate no-film and film-cooled baselines against reference data.",
            "Report convergence, mass imbalance, y+, mesh quality, and validation error metrics.",
        ]),
        ("7. Optimization experiment design", [
            "Compare Random/LHS, vanilla BO, zero-mean MFBO, piBO if implemented, and the proposed method.",
            "Use multiple random seeds and report mean plus dispersion.",
            "Run ablations for fixed prior, learnable prior, prior location, and wrong-prior robustness.",
        ]),
        ("8. Results", [
            "[RESULT: L1/L2 correlation and cost ratio.]",
            "[RESULT: validation plots.]",
            "[RESULT: BO convergence curves.]",
            "[RESULT: final design comparison.]",
            "[RESULT: ablation and wrong-prior robustness.]",
        ]),
        ("9. Discussion", [
            "Explain why the prior helps when it is imperfect.",
            "Discuss conditions under which multi-fidelity coupling helps or fails.",
            "State limitations: RANS fidelity, C3X-specific validation, prior dependence, and single-condition objective if robustness is not completed.",
        ]),
        ("10. Conclusions", [
            "Summarize method, validation, sample-efficiency result, and engineering implication.",
            "Use only numbers supported by final figures and tables.",
        ]),
    ]

    for heading, bullets in sections:
        add_paragraph(doc, heading, style="Heading 1")
        add_bullets(doc, bullets)

    add_paragraph(doc, "Planned figures and tables", style="Heading 1")
    add_table(
        doc,
        ["Item", "Purpose", "Status"],
        [
            ("Fig. 1", "Method workflow and data flow", "Can draft now"),
            ("Fig. 2", "C3X domain and film-row layout", "Can draft from current geometry"),
            ("Fig. 3", "No-film validation", "Needs final validation data"),
            ("Fig. 4", "Film-cooled baseline validation", "Needs film solve/data"),
            ("Fig. 5", "L1/L2 correlation and cost ratio", "Needs paired samples"),
            ("Fig. 6", "BO convergence versus equivalent L2 cost", "Needs BO runs"),
            ("Fig. 7", "Prior ablation and wrong-prior robustness", "Needs ablations"),
            ("Table 1", "Design variables and bounds", "Can draft now"),
            ("Table 2", "Fidelity hierarchy", "Can draft now"),
            ("Table 3", "CFD settings", "Partly ready"),
        ],
        [1300, 5600, 2460],
    )

    add_paragraph(doc, "CRediT authorship contribution statement", style="Heading 1")
    add_paragraph(doc, "[To be completed on the title page or manuscript according to final anonymization workflow: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Visualization, Writing - original draft, Writing - review and editing, Supervision, Funding acquisition.]")

    add_paragraph(doc, "Declaration of generative AI and AI-assisted technologies", style="Heading 1")
    add_paragraph(doc, "[If applicable, insert Elsevier-compliant statement before references. If AI tools were used only for grammar, spelling, or reference checking, no statement is required by the current policy. If used for manuscript organization or drafting assistance, declare the tool, purpose, human review, and author responsibility.]")

    add_paragraph(doc, "Declaration of competing interest", style="Heading 1")
    add_paragraph(doc, "[The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.]")

    add_paragraph(doc, "Funding", style="Heading 1")
    add_paragraph(doc, "[Insert funding statement. If none: This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.]")

    add_paragraph(doc, "Data availability", style="Heading 1")
    add_paragraph(doc, "[Insert repository DOI/link for code, processed CFD data, input layouts, and scripts. If some raw CFD data cannot be shared, state why and provide derived data sufficient to reproduce all figures.]")

    add_paragraph(doc, "References", style="Heading 1")
    add_paragraph(doc, "[1] Hylton et al., NASA CR-168015, no-film C3X validation reference.")
    add_paragraph(doc, "[2] Hylton et al., NASA CR-182133, film-cooled C3X validation reference.")
    add_paragraph(doc, "[3] Sellers, gaseous film cooling with multiple injection stations.")
    add_paragraph(doc, "[4] Kennedy and O'Hagan, multi-fidelity co-Kriging.")
    add_paragraph(doc, "[5] Balandat et al., BoTorch framework.")

    out = OUT_DIR / "AST_anonymized_manuscript_skeleton.docx"
    doc.save(out)
    return out


def build_title_page():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AST Title Page Template")

    add_paragraph(doc, "Article title", style="Heading 1")
    add_paragraph(doc, "Correlation-informed multi-fidelity Bayesian optimization for film-cooling hole layout on a C3X turbine vane")

    add_paragraph(doc, "Authors and affiliations", style="Heading 1")
    add_paragraph(doc, "[Author 1 given name family name]a, [Author 2 given name family name]b")
    add_paragraph(doc, "a [Department, Institution, Full postal address, Country, email]")
    add_paragraph(doc, "b [Department, Institution, Full postal address, Country, email]")

    add_paragraph(doc, "Corresponding author", style="Heading 1")
    add_table(
        doc,
        ["Field", "Content"],
        [
            ("Name", "[Corresponding author full name]"),
            ("Full postal address", "[Full address required by AST]"),
            ("Email", "[email@example.com]"),
            ("Phone", "[optional but useful for submission checklist]"),
        ],
        [2200, 7160],
    )

    add_paragraph(doc, "Acknowledgements", style="Heading 1")
    add_paragraph(doc, "[Include language editing, proof reading, technical assistance, compute access, and non-author support here only. Do not include acknowledgements in the anonymized manuscript.]")

    add_paragraph(doc, "Declaration of competing interests", style="Heading 1")
    add_paragraph(doc, "[Insert declaration or note that a separate declaration file will be uploaded.]")

    add_paragraph(doc, "Funding", style="Heading 1")
    add_paragraph(doc, "[Insert funder names and grant numbers, plus sponsor role. If none, use the standard no-specific-funding sentence.]")

    add_paragraph(doc, "Anonymization reminder", style="Heading 1")
    add_bullets(
        doc,
        [
            "Remove authors, affiliations, and acknowledgements from the main manuscript.",
            "Check document properties and file names before submission.",
            "Avoid self-identifying phrases such as 'in our previous work' unless anonymized.",
            "Keep CRediT details here until the journal requests non-anonymized files.",
        ],
    )

    out = OUT_DIR / "AST_title_page_template.docx"
    doc.save(out)
    return out


def build_highlights():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Highlights")

    add_paragraph(
        doc,
        "Highlights file for Aerospace Science and Technology. Each bullet is no more than 85 characters including spaces."
    )

    highlights = [
        "Film-cooling correlations define a Gaussian-process prior mean.",
        "Multi-fidelity BO jointly selects C3X designs and CFD fidelity.",
        "Learnable prior coefficients improve robustness to correlation bias.",
        "Cost-aware acquisition reduces equivalent high-fidelity CFD calls.",
        "C3X validation links optimization gains to turbine-vane cooling.",
    ]
    for item in highlights:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.add_run(f" [{len(item)} characters]").italic = True

    out = OUT_DIR / "AST_highlights.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    main = build_main_manuscript()
    title = build_title_page()
    highlights = build_highlights()
    print(main)
    print(title)
    print(highlights)
